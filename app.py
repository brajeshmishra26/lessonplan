from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
import sqlite3
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

app = Flask(__name__)

# Database setup
def init_db():
    conn = sqlite3.connect('lesson_plans.db')
    cursor = conn.cursor()
    
    # Create tables
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS lesson_plans (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            class_num INTEGER,
            subject TEXT,
            chapter_num INTEGER,
            chapter_title TEXT,
            month TEXT,
            year INTEGER,
            unit TEXT,
            learning_outcomes TEXT,
            teaching_points TEXT,
            teaching_aids TEXT,
            teaching_methodology TEXT,
            activity_planned TEXT,
            learning_skill_practice TEXT,
            life_skill_learnt TEXT,
            assignment_planned TEXT,
            assessment_planned TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Insert sample data (Class 5 Rimjhim Chapter 8)
    cursor.execute('''
        INSERT OR REPLACE INTO lesson_plans 
        (class_num, subject, chapter_num, chapter_title, month, year, unit,
         learning_outcomes, teaching_points, teaching_aids, teaching_methodology,
         activity_planned, learning_skill_practice, life_skill_learnt,
         assignment_planned, assessment_planned)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        5, 'Hindi (रिमझिम)', 8, 'वे दिन भी क्या दिन थे (विज्ञान कथा)', 'October', 2025, 'पाठ 8',
        '''1. छात्र विज्ञान कथा की विधा को समझ सकेंगे और उसकी विशेषताओं को पहचान सकेंगे
2. भविष्य की तकनीक और आविष्कारों के बारे में कल्पना कर सकेंगे
3. पाठ के माध्यम से भाषा की सुंदरता और शब्द भंडार में वृद्धि कर सकेंगे
4. कहानी के पात्रों के भावों और विचारों को समझ सकेंगे
5. वैज्ञानिक दृष्टिकोण विकसित कर तर्कसंगत चिंतन कर सकेंगे''',
        '''1. विज्ञान कथा का परिचय और इसकी विशेषताएं
2. कहानी के मुख्य पात्रों का चरित्र-चित्रण
3. भविष्य की तकनीक और वैज्ञानिक आविष्कारों की चर्चा
4. पाठ में आए नए शब्दों का अर्थ और प्रयोग
5. कहानी की मुख्य घटनाओं का क्रमवार विवरण
6. पात्रों के संवादों का भावपूर्ण वाचन
7. वैज्ञानिक सोच और कल्पनाशीलता का विकास
8. भाषा की शुद्धता और उच्चारण पर ध्यान
9. कहानी के माध्यम से मिलने वाले संदेश की व्याख्या
10. आधुनिक तकनीक और भविष्य की संभावनाओं पर चर्चा''',
        '''1. रिमझिम पाठ्य पुस्तक, श्यामपट्ट और चार्ट पेपर
2. विज्ञान और तकनीक से संबंधित चित्र और मॉडल
3. प्रोजेक्टर/स्मार्ट बोर्ड (यदि उपलब्ध हो)''',
        '''1. व्याख्यान विधि और प्रश्नोत्तर विधि का प्रयोग
2. समूहिक चर्चा और सहयोगी शिक्षण पद्धति''',
        'छात्रों को समूहों में बांटकर भविष्य के आविष्कारों पर चर्चा कराना और उन्हें अपनी कल्पना के आधार पर एक छोटी विज्ञान कथा लिखने के लिए प्रेरित करना',
        '''1. पठन कौशल - पाठ का शुद्ध उच्चारण और भावपूर्ण वाचन
2. लेखन कौशल - नए शब्दों का शुद्ध लेखन और वाक्य प्रयोग''',
        'वैज्ञानिक दृष्टिकोण और तर्कसंगत चिंतन का विकास, भविष्य की चुनौतियों के लिए तैयार होना',
        'पाठ के कठिन शब्दों के अर्थ लिखकर वाक्य प्रयोग करना तथा अपनी कल्पना के आधार पर 10 पंक्तियों में एक छोटी विज्ञान कथा लिखना',
        '''1. मौखिक प्रश्नों के द्वारा पाठ की समझ का आकलन
2. शब्द भंडार और भाषा प्रयोग की जांच
3. रचनात्मक लेखन और कल्पनाशीलता का मूल्यांकन
4. सामूहिक गतिविधि में भागीदारी और सहयोग का आकलन
5. गृहकार्य की जांच और व्यक्तिगत प्रगति का मूल्यांकन'''
    ))
    
    # Add more sample lesson plans for demonstration
    sample_lessons = [
        # Class 4 lessons
        (4, 'Hindi (रिमझिम)', 5, 'नन्हा फनकार', 'October', 2025, 'पाठ 5',
         '''1. छात्र कहानी के माध्यम से कला और रचनात्मकता को समझ सकेंगे
2. पात्रों के चरित्र और उनकी विशेषताओं को पहचान सकेंगे
3. कहानी के संदेश को समझकर अपने जीवन में लागू कर सकेंगे
4. नए शब्दों का अर्थ समझकर शब्द भंडार बढ़ा सकेंगे
5. कहानी के माध्यम से भावनाओं की अभिव्यक्ति कर सकेंगे''',
         '''1. कहानी का परिचय और मुख्य पात्रों की चर्चा
2. नन्हा फनकार के चरित्र की विशेषताएं
3. कहानी में आए नए शब्दों का अर्थ और प्रयोग
4. कहानी की मुख्य घटनाओं का विवरण
5. पात्रों के संवादों का अभ्यास
6. कहानी के संदेश की व्याख्या
7. रचनात्मकता और कला के महत्व पर चर्चा
8. भाषा की शुद्धता पर ध्यान
9. कहानी से मिली सीख पर चर्चा
10. व्यावहारिक जीवन में कहानी का प्रयोग''',
         '''1. रिमझिम पाठ्य पुस्तक और चार्ट पेपर
2. कहानी से संबंधित चित्र और पोस्टर
3. श्यामपट्ट और मार्कर''',
         '''1. कहानी सुनाने की विधि और प्रश्नोत्तर
2. समूहिक चर्चा और रोल प्ले''',
         'छात्रों से कहानी के पात्रों का चरित्र चित्रण करवाना और अपनी रचनात्मकता दिखाने के लिए कहना',
         '''1. सुनना और समझना - कहानी को ध्यान से सुनना
2. बोलना - पात्रों के संवाद बोलना और अपनी बात कहना''',
         'रचनात्मकता और कला के प्रति सकारात्मक दृष्टिकोण विकसित करना',
         'कहानी के मुख्य पात्रों के बारे में 5 वाक्य लिखना और नए शब्दों के वाक्य प्रयोग करना',
         '''1. कहानी की समझ का मौखिक आकलन
2. नए शब्दों की जानकारी की जांच
3. पात्रों के चरित्र की समझ का मूल्यांकन
4. गृहकार्य की जांच और फीडबैक'''),
        
        # Class 6 lessons
        (6, 'Hindi (वसंत)', 3, 'नादान दोस्त', 'September', 2025, 'पाठ 3',
         '''1. छात्र कहानी के माध्यम से मित्रता के महत्व को समझ सकेंगे
2. गलत और सही में अंतर कर सकेंगे
3. कहानी के पात्रों की मानसिकता को समझ सकेंगे
4. भाषा की सुंदरता और अभिव्यक्ति में सुधार कर सकेंगे
5. नैतिक मूल्यों को समझकर अपनाने की प्रेरणा ले सकेंगे''',
         '''1. कहानी का परिचय और लेखक की जानकारी
2. मुख्य पात्रों का परिचय और चरित्र चित्रण
3. कहानी की मुख्य घटनाओं का विवरण
4. संवादों का अभ्यास और भावपूर्ण वाचन
5. कहानी में आए कठिन शब्दों का अर्थ
6. कहानी के संदेश और शिक्षा पर चर्चा
7. मित्रता के सच्चे अर्थ की व्याख्या
8. भाषा की शुद्धता और व्याकरण
9. पात्रों के व्यवहार का विश्लेषण
10. कहानी से जुड़े जीवन मूल्यों की चर्चा''',
         '''1. वसंत पाठ्य पुस्तक और संदर्भ सामग्री
2. कहानी से संबंधित चित्र और मॉडल
3. श्यामपट्ट और प्रोजेक्टर''',
         '''1. कहानी विधि और चर्चा विधि
2. प्रश्नोत्तर और समूहिक गतिविधि''',
         'छात्रों को समूहों में बांटकर कहानी का नाटकीकरण करवाना और मित्रता पर चर्चा',
         '''1. पठन कौशल - शुद्ध उच्चारण और भावपूर्ण वाचन
2. लेखन कौशल - कहानी का सार लेखन''',
         'सच्ची मित्रता और नैतिक मूल्यों की समझ विकसित करना',
         'कहानी का सार 10 वाक्यों में लिखना और मित्रता पर अपने विचार व्यक्त करना',
         '''1. कहानी की समझ का मौखिक मूल्यांकन
2. चरित्र चित्रण की समझ की जांच
3. भाषा प्रयोग और शब्द भंडार का आकलन
4. नैतिक मूल्यों की समझ की जांच'''),
         
        # Class 3 lessons  
        (3, 'Hindi (रिमझिम)', 7, 'हमारे जवान', 'November', 2025, 'पाठ 7',
         '''1. छात्र देश के वीर जवानों के बारे में जान सकेंगे
2. देशभक्ति की भावना विकसित कर सकेंगे
3. जवानों के साहस और वीरता की कहानी समझ सकेंगे
4. कविता के भाव और संदेश को समझ सकेंगे
5. उचित उच्चारण और लय के साथ कविता पढ़ सकेंगे''',
         '''1. कविता का परिचय और कवि की जानकारी
2. जवानों की वीरता और साहस की चर्चा
3. कविता में आए नए शब्दों का अर्थ
4. कविता का भावपूर्ण वाचन
5. देशभक्ति की भावना पर चर्चा
6. कविता की लय और छंद की समझ
7. जवानों के जीवन पर चर्चा
8. कविता के माध्यम से संदेश की व्याख्या
9. भाषा की सुंदरता पर ध्यान
10. देश सेवा के महत्व पर चर्चा''',
         '''1. रिमझिम पाठ्य पुस्तक और कविता चार्ट
2. जवानों के चित्र और तिरंगा
3. श्यामपट्ट और संगीत वाद्य यंत्र''',
         '''1. कविता विधि और संवादात्मक शिक्षण
2. समूहिक गायन और अभिनय विधि''',
         'छात्रों के साथ देशभक्ति गीत गाना और जवानों के बारे में चर्चा करना',
         '''1. सुनना और बोलना - कविता सुनना और सुनाना
2. पढ़ना - लयबद्ध और भावपूर्ण वाचन''',
         'देशभक्ति और वीरता की भावना का विकास',
         'कविता के मुख्य भाव को अपने शब्दों में लिखना और जवानों के बारे में 3 वाक्य लिखना',
         '''1. कविता की समझ का मौखिक आकलन
2. उच्चारण और लय की जांच
3. देशभक्ति की भावना का मूल्यांकन
4. शब्द ज्ञान की जांच''')
    ]
    
    # Insert sample lessons
    for lesson in sample_lessons:
        cursor.execute('''
            INSERT OR REPLACE INTO lesson_plans 
            (class_num, subject, chapter_num, chapter_title, month, year, unit,
             learning_outcomes, teaching_points, teaching_aids, teaching_methodology,
             activity_planned, learning_skill_practice, life_skill_learnt,
             assignment_planned, assessment_planned)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', lesson)
    
    conn.commit()
    conn.close()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/test_print')
def test_print():
    return '''
    <!DOCTYPE html>
    <html>
    <head>
        <title>Print Test</title>
        <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
        <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
        <style>
            @media print { .no-print { display: none !important; } }
        </style>
    </head>
    <body>
        <div class="container mt-5">
            <div class="no-print text-center mb-4">
                <button id="printBtn1" class="btn btn-success me-2">
                    <i class="fas fa-print"></i> Test Print
                </button>
                <button onclick="window.print()" class="btn btn-primary me-2">
                    <i class="fas fa-print"></i> Direct Print
                </button>
            </div>
            <h2>Print Test Content</h2>
            <p>This is test content for print functionality.</p>
        </div>
        <script>
            document.addEventListener('DOMContentLoaded', function() {
                document.getElementById('printBtn1').addEventListener('click', function() {
                    console.log('Print button clicked');
                    window.print();
                });
            });
        </script>
    </body>
    </html>
    '''

@app.route('/get_chapters/<int:class_num>')
def get_chapters(class_num):
    conn = sqlite3.connect('lesson_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT DISTINCT chapter_num, chapter_title 
        FROM lesson_plans 
        WHERE class_num = ? 
        ORDER BY chapter_num
    ''', (class_num,))
    
    chapters = cursor.fetchall()
    conn.close()
    
    return jsonify([{'chapter_num': ch[0], 'chapter_title': ch[1]} for ch in chapters])

@app.route('/get_lesson_plans')
@app.route('/get_lesson_plans/<int:class_num>')
def get_lesson_plans(class_num=None):
    conn = sqlite3.connect('lesson_plans.db')
    cursor = conn.cursor()
    
    if class_num:
        cursor.execute('''
            SELECT class_num, subject, chapter_num, chapter_title, month, year
            FROM lesson_plans 
            WHERE class_num = ?
            ORDER BY chapter_num
        ''', (class_num,))
    else:
        cursor.execute('''
            SELECT class_num, subject, chapter_num, chapter_title, month, year
            FROM lesson_plans 
            ORDER BY class_num, chapter_num
        ''')
    
    lesson_plans = cursor.fetchall()
    conn.close()
    
    return jsonify([{
        'class_num': lp[0],
        'subject': lp[1], 
        'chapter_num': lp[2],
        'chapter_title': lp[3],
        'month': lp[4],
        'year': lp[5]
    } for lp in lesson_plans])

@app.route('/lesson_plan/<int:class_num>/<int:chapter_num>')
def view_lesson_plan(class_num, chapter_num):
    conn = sqlite3.connect('lesson_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM lesson_plans 
        WHERE class_num = ? AND chapter_num = ?
    ''', (class_num, chapter_num))
    
    lesson_plan = cursor.fetchone()
    conn.close()
    
    if not lesson_plan:
        return "Lesson plan not found!", 404
    
    # Convert tuple to dictionary for easy template access
    columns = ['id', 'class_num', 'subject', 'chapter_num', 'chapter_title', 
               'month', 'year', 'unit', 'learning_outcomes', 'teaching_points',
               'teaching_aids', 'teaching_methodology', 'activity_planned',
               'learning_skill_practice', 'life_skill_learnt', 'assignment_planned',
               'assessment_planned', 'created_at']
    
    lesson_data = dict(zip(columns, lesson_plan))
    
    return render_template('lesson_plan.html', lesson=lesson_data)

@app.route('/print/<int:class_num>/<int:chapter_num>')
def print_lesson_plan(class_num, chapter_num):
    conn = sqlite3.connect('lesson_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM lesson_plans 
        WHERE class_num = ? AND chapter_num = ?
    ''', (class_num, chapter_num))
    
    lesson_plan = cursor.fetchone()
    conn.close()
    
    if not lesson_plan:
        return "Lesson plan not found!", 404
    
    # Convert tuple to dictionary for easy template access
    columns = ['id', 'class_num', 'subject', 'chapter_num', 'chapter_title', 
               'month', 'year', 'unit', 'learning_outcomes', 'teaching_points',
               'teaching_aids', 'teaching_methodology', 'activity_planned',
               'learning_skill_practice', 'life_skill_learnt', 'assignment_planned',
               'assessment_planned', 'created_at']
    
    lesson_data = dict(zip(columns, lesson_plan))
    
    # Return a print-optimized version
    return render_template('print_lesson_plan.html', lesson=lesson_data)

@app.route('/download/<int:class_num>/<int:chapter_num>')
def download_lesson_plan(class_num, chapter_num):
    conn = sqlite3.connect('lesson_plans.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM lesson_plans 
        WHERE class_num = ? AND chapter_num = ?
    ''', (class_num, chapter_num))
    
    lesson_plan = cursor.fetchone()
    conn.close()
    
    if not lesson_plan:
        return "Lesson plan not found!", 404
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading('LESSON PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Basic Information
    basic_info_para = doc.add_paragraph()
    basic_info_para.add_run(f'Class: {lesson_plan[1]}    ')
    basic_info_para.add_run(f'Month: {lesson_plan[5]}    ')
    basic_info_para.add_run(f'Year: {lesson_plan[6]}\n')
    basic_info_para.add_run(f'Subject: {lesson_plan[2]}    ')
    basic_info_para.add_run(f'Unit: {lesson_plan[7]}\n')
    basic_info_para.add_run(f'Topic: {lesson_plan[4]}')
    
    doc.add_paragraph()
    
    # Add all sections
    sections = [
        ('Learning Outcomes (शिक्षण उद्देश्य)', lesson_plan[8]),
        ('Teaching Points (शिक्षण बिंदु)', lesson_plan[9]),
        ('Teaching Aids (शिक्षण सहायक सामग्री)', lesson_plan[10]),
        ('Teaching Methodology (शिक्षण विधि)', lesson_plan[11]),
        ('Activity Planned (नियोजित गतिविधि)', lesson_plan[12]),
        ('Learning Skill Practice (अधिगम कौशल अभ्यास)', lesson_plan[13]),
        ('Life Skill Learnt (जीवन कौशल)', lesson_plan[14]),
        ('Assignment Planned (गृहकार्य)', lesson_plan[15]),
        ('Assessment Planned (मूल्यांकन योजना)', lesson_plan[16])
    ]
    
    for section_title, content in sections:
        doc.add_heading(section_title, level=2)
        
        # Split content into lines and add as bullet points
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line.strip())
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.add_run('तैयारकर्ता: [शिक्षक का नाम]\n')
    footer_para.add_run(f'दिनांक: {datetime.now().strftime("%d %B, %Y")}')
    
    # Save to memory
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    filename = f'lesson_plan_class{class_num}_chapter{chapter_num}.docx'
    
    return send_file(
        doc_io,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    init_db()
    app.run(debug=True, host='0.0.0.0', port=5000)